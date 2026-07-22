import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.services.document_parser import extract_document_text


def test_extract_txt(tmp_path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Hello enterprise world.", encoding="utf-8")

    pages = extract_document_text(file_path)

    assert pages == [{"page": 1, "text": "Hello enterprise world."}]


def test_extract_docx(tmp_path):
    file_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(file_path)

    pages = extract_document_text(file_path)

    assert len(pages) == 1
    assert "First paragraph." in pages[0]["text"]
    assert "Second paragraph." in pages[0]["text"]


def test_extract_pdf_with_no_text_returns_no_pages(tmp_path):
    file_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(file_path, "wb") as handle:
        writer.write(handle)

    pages = extract_document_text(file_path)

    assert pages == []


def test_extract_unsupported_type_raises_value_error(tmp_path):
    file_path = tmp_path / "sample.csv"
    file_path.write_text("a,b,c", encoding="utf-8")

    with pytest.raises(ValueError):
        extract_document_text(file_path)


def test_extract_is_case_insensitive_on_extension(tmp_path):
    file_path = tmp_path / "sample.TXT"
    file_path.write_text("Uppercase extension still works.", encoding="utf-8")

    pages = extract_document_text(file_path)

    assert pages == [{"page": 1, "text": "Uppercase extension still works."}]


def test_extract_txt_preserves_unicode_content(tmp_path):
    file_path = tmp_path / "sample.txt"
    text = "Café policy: employees may expense up to €20 — no exceptions."
    file_path.write_text(text, encoding="utf-8")

    pages = extract_document_text(file_path)

    assert pages == [{"page": 1, "text": text}]


def test_extract_docx_skips_whitespace_only_paragraphs(tmp_path):
    file_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Real content.")
    document.add_paragraph("   ")
    document.add_paragraph("")
    document.save(file_path)

    pages = extract_document_text(file_path)

    assert pages[0]["text"] == "Real content."
